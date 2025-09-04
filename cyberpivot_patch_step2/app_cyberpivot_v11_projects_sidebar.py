# -*- coding: utf-8 -*-
import os, io, zipfile, math
from datetime import datetime

import streamlit as st
import pandas as pd

# Matplotlib (headless)
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Modules locaux (je te fournis les fichiers après)
import auth
import storage
import errors
import standards
import catalog

# -----------------------------
# Helpers
# -----------------------------
# --- Helpers robustesse ---
def safe_upsert(audit_id, payload):
    try:
        storage.upsert_response(audit_id, payload)
    except Exception as e:
        errors.report_error('Sauvegarde de la réponse', e)

def safe_rerun():
    if not st.session_state.get("_just_rerun"):
        st.session_state._just_rerun = True
        st.rerun()
    else:
        st.session_state._just_rerun = False

def normalize_text(s):
    if pd.isna(s):
        return ""
    return str(s).strip()

def score_global(df):
    if df.empty:
        return 0.0
    return round(df["score"].sum() / (2 * len(df)) * 100, 2)

def radar_or_bars(scores_series):
    labels = list(scores_series.index)
    values = list(scores_series.values)
    if len(labels) == 0:
        return None
    fig = None
    if len(labels) < 3:
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.bar(labels, values)
        ax.set_ylim(0, 100)
        ax.set_ylabel("Taux de conformité (%)")
        ax.set_title("Conformité par domaine")
    else:
        angles = [n / float(len(labels)) * 2 * math.pi for n in range(len(labels))]
        values_c = values + values[:1]
        angles_c = angles + angles[:1]
        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
        ax.plot(angles_c, values_c, linewidth=1, linestyle="solid")
        ax.fill(angles_c, values_c, alpha=0.35)
        ax.set_xticks(angles)
        ax.set_xticklabels(labels)
        ax.set_ylim(0, 100)
        ax.set_title("Radar de conformité par domaine")
    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()

LEGENDE = {
    "Conforme": "Toutes les exigences sont satisfaites avec preuves suffisantes et datées.",
    "Partiellement conforme": "Exigences globalement adressées mais preuves incomplètes / contrôles partiels.",
    "Non conforme": "Exigence non traitée ou preuve absente/obsolète.",
    "À définir": "Information non collectée (en attente d’entretien/preuve)."
}
NIVEAUX = ["Conforme", "Partiellement conforme", "Non conforme", "À définir"]
SCORE_MAP = {"Conforme": 2, "Partiellement conforme": 1, "Non conforme": 0, "À définir": 0}

MODELES_DOCS = {
    "Politique IA": "templates/politique_IA.docx",
    "Méthode risques IA": "templates/methode_gestion_risques.docx",
    "Analyse risques IA": "templates/modele_registre_risques.xlsx",
    "Traitement des risques IA": "templates/plan_traitement_risques.docx",
    "Usage licite de l'IA": "templates/checklist_conformite_juridique.docx",
    "Gestion des biais IA": "templates/protocole_tests_biais.docx",
    "Protection des données d’entraînement": "templates/politique_securite_donnees.docx",
    "Protection des modèles IA": "templates/standard_durcissement_modeles.docx",
    "Mise en production sécurisée": "templates/procedure_cicd_securisee.docx",
    "Journalisation IA": "templates/politique_journalisation.docx",
    "Supervision IA": "templates/plan_supervision_kpis.docx",
    "Base légale de traitement": "templates/modele_dpia.docx",
    "Minimisation des données": "templates/standard_minimisation_donnees.docx",
    "Comité IA": "templates/charte_comite_IA.docx",
    "Responsabilités IA": "templates/modele_fiche_responsable_IA.docx",
}
os.makedirs("templates", exist_ok=True)

def isaca_word_report(df, client_name, auditor_name, taux_global, radar_png):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("Rapport d’audit — ISO/IEC 42001", 0)
    doc.add_paragraph(f"Client : {client_name}")
    doc.add_paragraph(f"Auditeur : {auditor_name}")
    doc.add_paragraph("Date : {}".format(datetime.now().strftime("%d/%m/%Y")))
    doc.add_paragraph("Classification : Diffusion restreinte – Document confidentiel")
    doc.add_page_break()

    doc.add_heading("Sommaire", level=1)
    doc.add_paragraph("Le sommaire sera mis à jour à l’ouverture du document (Références > Mettre à jour la table).")
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Ce rapport présente les résultats de l’audit selon la norme ISO/IEC 42001, structuré au format ISACA."
    )
    doc.add_paragraph("Objectifs :")
    doc.add_paragraph("• Évaluer le niveau de maturité et de conformité ;")
    doc.add_paragraph("• Identifier les non-conformités et proposer des recommandations réalistes ;")
    doc.add_paragraph("• Fournir une synthèse des risques et un plan d’action priorisé.")

    doc.add_page_break()
    doc.add_heading("2. Constatations d’audit", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Domaine"
    hdr[1].text = "ID"
    hdr[2].text = "Point de contrôle"
    hdr[3].text = "Constat / Critère"
    hdr[4].text = "Recommandation"

    subset = df[df["score"] < 2].copy()
    for _, r in subset.iterrows():
        row = table.add_row().cells
        row[0].text = str(r["domaine"])
        row[1].text = str(r["id"])
        row[2].text = str(r["item"])
        row[3].text = str(r.get("critere", ""))
        row[4].text = str(r.get("recommandation", ""))

    doc.add_page_break()
    doc.add_heading("3. Synthèse des risques", level=1)
    doc.add_paragraph("Taux global de conformité : {} %.".format(taux_global))

    if radar_png:
        doc.add_paragraph("Radar de conformité par domaine :")
        pic = doc.add_picture(io.BytesIO(radar_png), width=Inches(5.8))
        last_par = doc.paragraphs[-1]
        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4. Conclusion", level=1)
    doc.add_paragraph(
        "Au vu des constats présentés, des améliorations sont recommandées pour élever le niveau de conformité et la maîtrise des risques."
    )

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def build_zip_dossier_audit(df, client_name, auditor_name):
    df_repondu = df[df["niveau"] != "À définir"].copy()
    taux = score_global(df_repondu)
    spd = (df_repondu.groupby("domaine")["score"].mean() * 50) if not df_repondu.empty else pd.Series(dtype=float)
    radar_bytes = radar_or_bars(spd) if not df_repondu.empty else None
    report_bytes = isaca_word_report(df_repondu, client_name, auditor_name, taux, radar_bytes)

    zbio = io.BytesIO()
    with zipfile.ZipFile(zbio, mode="w", compression=zipfile.ZIP_DEFLATED) as z:
        # Excel des réponses
        xbio = io.BytesIO()
        df.to_excel(xbio, index=False, sheet_name="Audit")
        z.writestr("Audit_CyberPivot.xlsx", xbio.getvalue())
        # Rapport Word
        z.writestr("Rapport_ISO42001_ISACA_CyberPivot.docx", report_bytes)
        # Modèles manquants (notes)
        subset = df_repondu[df_repondu["score"] < 2]
        if not subset.empty:
            for _, r in subset.iterrows():
                item = str(r["item"])
                modele = MODELES_DOCS.get(item)
                if modele and os.path.isfile(modele):
                    with open(modele, "rb") as fh:
                        z.writestr("MODELES/{}".format(os.path.basename(modele)), fh.read())
                else:
                    note = "Pas de modèle mappé ou fichier introuvable pour l’item: {}".format(item)
                    z.writestr("MODELES/{}_README.txt".format(item.replace(" ", "_")), note.encode("utf-8"))
    zbio.seek(0)
    return zbio.getvalue()

# -----------------------------
# CONFIG APP
# -----------------------------
st.set_page_config(page_title="CyberPivot™ – Audit ISACA ISO 42001", layout="wide")
st.title("🧠 CyberPivot™ — Audit dynamique ISO 42001 (format ISACA)")

# Init DBs
auth.init_db()
storage.init_db()
try:
    storage.migrate_add_yaml_columns()
except Exception:
    pass

# -----------------------------
# SESSION / LOGIN
# -----------------------------
if "user" not in st.session_state:
    st.session_state.user = None

with st.sidebar:
    st.markdown("### 👤 Compte")
    if st.session_state.user:
        name = st.session_state.user.get('name') if isinstance(st.session_state.user, dict) else str(st.session_state.user)
        role = st.session_state.user.get('role') if isinstance(st.session_state.user, dict) else "auditeur"
        st.write(f"Connecté : **{name}** ({role})")
        if st.button("Se déconnecter", key="logout_btn"):
            st.session_state.user = None
            safe_rerun()
    else:
        tab_login, tab_signup = st.tabs(["Se connecter", "Créer un admin (init)"])
        with tab_login:
            email = st.text_input("Email", key="login_email")
            pwd = st.text_input("Mot de passe", type="password", key="login_pwd")
            if st.button("Connexion", key="login_btn"):
                user, err = auth.authenticate(email, pwd)
                if err:
                    st.error(err)
                else:
                    st.session_state.user = user
                    st.success("Connexion réussie.")
                    safe_rerun()
        with tab_signup:
            st.caption("Si c’est le premier démarrage et qu’il n’y a aucun compte, créez un **administrateur**.")
            admin_email = st.text_input("Email administrateur", key="adm_email")
            admin_name = st.text_input("Nom complet", key="adm_name")
            admin_pwd = st.text_input("Mot de passe", type="password", key="adm_pwd")
            if st.button("Créer l’admin", key="adm_create_btn"):
                ok, msg = auth.create_user(admin_email, admin_name, "admin", admin_pwd)
                st.success(msg) if ok else st.error(msg)

    if not st.session_state.user:
        st.stop()

    # ----- Norme -----
    st.markdown("### 📚 Norme")
    try:
        std_list = standards.list_standards()
    except Exception:
        std_list = []
    if not std_list:
        std_list = ["iso42001"]
    std_name = st.selectbox("Norme active", std_list, key="std_name")
    #st.session_state.std_name = std_name

    # ----- Projets & Audits -----
    st.markdown("---")
    st.markdown("### 📁 Projets")
    try:
        _projects = storage.list_projects()
    except Exception as _e:
        _projects = []
        st.warning(f"Impossible de lister les projets: {_e}")
    proj_labels = [f"{p['id']} — {p['name']} ({p.get('client','')})" for p in _projects] or ["Aucun"]
    default_proj_idx = 0
    if "current_project_id" in st.session_state and _projects:
        for idx, p in enumerate(_projects):
            if p["id"] == st.session_state["current_project_id"]:
                default_proj_idx = idx
                break
    selected_proj = st.selectbox("Projet courant", proj_labels, index=default_proj_idx, key="proj_select")
    if _projects:
        try:
            st.session_state.current_project_id = int(selected_proj.split(" — ")[0])
        except Exception:
            pass
    else:
        if st.button("Créer un projet par défaut", key="create_default_proj_btn"):
            pid = storage.create_project("Default Project", "Default Client")
            st.session_state.current_project_id = pid
            st.success(f"Projet créé (id={pid}).")
            safe_rerun()

    with st.expander("➕ Créer un projet", expanded=False):
        np_name = st.text_input("Nom du projet", key="np_name")
        np_client = st.text_input("Client", key="np_client")
        if st.button("Créer le projet", key="btn_create_project"):
            if np_name.strip():
                pid = storage.create_project(np_name.strip(), np_client.strip())
                st.session_state.current_project_id = pid
                st.success(f"Projet créé (id={pid}).")
                safe_rerun()
            else:
                st.error("Le nom du projet est requis.")

    st.markdown("### 🗂️ Audits")
    _audits = storage.list_audits(project_id=st.session_state.get("current_project_id")) if st.session_state.get("current_project_id") else []
    aud_labels = [f"{a['id']} — {a['standard']} {a.get('version','')}" for a in _audits] or ["Aucun"]
    default_aud_idx = 0
    if "current_audit_id" in st.session_state and _audits:
        for idx, a in enumerate(_audits):
            if a["id"] == st.session_state["current_audit_id"]:
                default_aud_idx = idx
                break
    selected_aud = st.selectbox("Audit courant", aud_labels, index=default_aud_idx, key="aud_select")
    if _audits:
        try:
            st.session_state.current_audit_id = int(selected_aud.split(" — ")[0])
        except Exception:
            pass

    col_a1, col_a2 = st.columns(2)
    with col_a1:
        if st.button("➕ Nouvel audit (norme sélectionnée)", key="btn_new_audit"):
            try:
                meta = standards.load_standard(st.session_state.get("std_name","iso42001"))
                version = meta.get("version","")
            except Exception:
                version = ""
            aid = storage.create_audit(st.session_state.current_project_id, st.session_state.get("std_name","iso42001"), version)
            st.session_state.current_audit_id = aid
            st.success(f"Audit créé (id={aid}).")
            safe_rerun()
    with col_a2:
        if st.button("🗑️ Supprimer l’audit courant", key="btn_delete_audit"):
            if st.session_state.get("current_audit_id"):
                storage.delete_audit(st.session_state["current_audit_id"])
                st.success("Audit supprimé.")
                st.session_state.current_audit_id = None
                safe_rerun()
            else:
                st.info("Aucun audit sélectionné.")

# -----------------------------
# Sidebar — Paramètres du rapport
# -----------------------------
st.sidebar.header("⚙️ Paramètres du rapport")
client_name = st.sidebar.text_input("Nom du client", "Organisation Auditée", key="client_name")
auditor_name = st.sidebar.text_input("Auditeur", "Équipe CyberPivot™", key="auditor_name")
st.sidebar.markdown("### ℹ️ Légende des niveaux")
for k, v in LEGENDE.items():
    st.sidebar.markdown(f"**{k}** — {v}")

# -----------------------------
# UPLOAD / LOAD CATALOG
# -----------------------------
st.subheader("📥 Import du questionnaire (Excel standardisé)")
st.caption("Choisis une norme ou importe un Excel pour générer un YAML mis en cache.")
uploaded = st.file_uploader(
    "Fichier .xlsx (colonnes : Domaine, ID, Item, Question, Objectif, Preuve attendue, Référence, Critère, Recommandation)",
    type=["xlsx"], key="xlsx_import"
)

df_src = None
audit_id = st.session_state.get("current_audit_id")

# Si Excel importé → générer YAML et lier à l’audit
if uploaded is not None and audit_id:
    try:
        df_x = pd.read_excel(uploaded, engine='openpyxl')
        df_x.columns = [str(c).strip() for c in df_x.columns]
        ypath = catalog.cache_catalog_from_excel(
            audit_id, df_x,
            title=f"Audit {st.session_state.get('std_name','')} — {client_name}",
            version=datetime.now().strftime("%Y%m%d")
        )
        st.success(f"Questionnaire importé et mis en cache: {ypath}")
    except Exception as _e:
        st.error(f"Erreur d'import Excel: {_e}")

# Charger le DF : priorité au YAML caché lié à l’audit, sinon YAML de la norme
try:
    audits = storage.list_audits(project_id=st.session_state.get('current_project_id'))
    audit = next((a for a in audits if a['id'] == audit_id), None)
    if audit is not None:
        df_src = catalog.load_catalog_for_audit(audit)
except Exception:
    df_src = None

if df_src is None:
    try:
        std_meta = standards.load_standard(st.session_state.get('std_name','iso42001'))
        df_src = standards.flatten_to_dataframe(std_meta)
        st.info("Chargé depuis la norme YAML (pas d'Excel).")
    except Exception as _e:
        st.error(f"Impossible de charger la norme: {_e}")
        df_src = pd.DataFrame()

EXPECTED = ["Domaine", "ID", "Item", "Question", "Objectif", "Preuve attendue", "Référence", "Critère", "Recommandation"]
missing = [c for c in EXPECTED if c not in df_src.columns]
if missing or df_src.empty:
    st.info("Choisis une norme ou importe un Excel pour démarrer.")
    st.stop()

# -----------------------------
# FORMULAIRE DYNAMIQUE + PERSISTENCE
# -----------------------------
st.subheader("📝 Questionnaire")

df_src["Domaine"] = df_src["Domaine"].replace("", pd.NA).ffill().fillna("")
domaines_uniques = sorted(df_src["Domaine"].dropna().unique().tolist())

# Préremplissage
existing = { (r["qid"] or r["item"]): r for r in storage.get_responses(audit_id) } if audit_id else {}

reponses = []

for d in domaines_uniques:
    bloc = df_src[df_src["Domaine"] == d].reset_index(drop=True)
    with st.expander(f"📂 {d}", expanded=False):
        page_size = 12
        total = len(bloc)
        pages = (total + page_size - 1) // page_size
        p = st.number_input(f"Page (1-{pages}) – {total} questions", min_value=1, max_value=max(1, pages), value=1, step=1, key=f"pg_{d}")
        start = (p - 1) * page_size
        end = min(start + page_size, total)
        view = bloc.iloc[start:end]

        for _, r in view.iterrows():
            qid = normalize_text(r["ID"]) or normalize_text(r["Item"])
            pref = existing.get(qid, {})
            col1, col2 = st.columns([3, 2])
            with col1:
                st.markdown(f"**{qid} – {normalize_text(r['Item'])}**")
                st.caption(normalize_text(r["Question"]))
                if normalize_text(r["Objectif"]):
                    st.write("🎯 Objectif :", normalize_text(r["Objectif"]))
                if normalize_text(r["Preuve attendue"]):
                    st.write("📎 Preuve attendue :", normalize_text(r["Preuve attendue"]))
                if normalize_text(r["Référence"]):
                    st.write("🔖 Référence :", normalize_text(r["Référence"]))
                nv_idx = NIVEAUX.index(pref.get("level")) if pref.get("level") in NIVEAUX else 3
                niveau = st.selectbox("Niveau de conformité", NIVEAUX, index=nv_idx, key=f"conf_{audit_id}_{qid}")
                commentaire = st.text_area("Commentaire (optionnel)", value=pref.get("comment",""), key=f"com_{audit_id}_{qid}", placeholder="Observations / précisions…")
            with col2:
                files = st.file_uploader("Joindre des preuves (multi)", type=["pdf","docx","xlsx","png","jpg","jpeg"], key=f"up_{audit_id}_{qid}", accept_multiple_files=True)
                names = [f.name for f in files] if files else []

            rec = {
                "domaine": d,
                "id": qid,
                "item": normalize_text(r["Item"]),
                "question": normalize_text(r["Question"]),
                "niveau": niveau,
                "score": SCORE_MAP.get(niveau, 0),
                "critere": normalize_text(r["Critère"]),
                "recommandation": normalize_text(r["Recommandation"]),
                "preuves": names,
                "commentaire": commentaire
            }
            reponses.append(rec)

            # Sauvegarde immédiate
            safe_upsert(audit_id, {
                "domain": d,
                "qid": qid,
                "item": rec["item"],
                "question": rec["question"],
                "level": rec["niveau"],
                "score": rec["score"],
                "criterion": rec["critere"],
                "recommendation": rec["recommandation"],
                "comment": rec["commentaire"],
                "evidence": names,
            })

df_form = pd.DataFrame(reponses)
df_repondu = df_form[df_form["niveau"] != "À définir"].copy()

st.sidebar.metric("Taux global de conformité", f"{score_global(df_repondu) if not df_repondu.empty else 0} %")

# -----------------------------
# SYNTHÈSE
# -----------------------------
st.divider()
st.subheader("📊 Synthèse")

taux_glob = score_global(df_repondu)
c1, c2 = st.columns([1, 2])
with c1:
    st.metric("Taux global de conformité", "{} %".format(taux_glob))
    st.write("Par domaine :")
    if not df_repondu.empty:
        spd = (df_repondu.groupby("domaine")["score"].mean() * 50).round(2)
        st.dataframe(spd.reset_index().rename(columns={"score": "Taux (%)"}), use_container_width=True)
    else:
        st.info("Aucune réponse utile.")
with c2:
    st.write("Radar / Barres")
    if not df_repondu.empty:
        spd = (df_repondu.groupby("domaine")["score"].mean() * 50)
        img = radar_or_bars(spd)
        if img:
            st.image(img, caption="Conformité par domaine", use_column_width=True)
    else:
        st.info("Le graphique s’affichera quand il y aura des réponses (hors 'À définir').")

st.write("Non-conformités & partiels :")
subset = df_repondu[df_repondu["score"] < 2]
if subset.empty:
    st.success("Aucune non-conformité détectée.")
else:
    for _, row in subset.iterrows():
        st.warning("🔎 [{}] {} → {}\n💡 Recommandation : {}".format(
            row["domaine"], row["item"], row["niveau"], row.get("recommandation", "")
        ))

# -----------------------------
# EXPORT
# -----------------------------
st.divider()
st.subheader("📤 Export (rapport ISACA + modèles)")

if st.button("📥 Télécharger le dossier d’audit", key="export_zip_btn"):
    if df_repondu.empty:
        st.error("Veuillez renseigner au moins une réponse (≠ 'À définir') avant d’exporter.")
    else:
        zip_bytes = build_zip_dossier_audit(df_form, client_name, auditor_name)
        st.success("✅ Dossier généré.")
        if not subset.empty:
            missing_models = [it for it in subset['item'].unique() if not MODELES_DOCS.get(it) or not os.path.isfile(MODELES_DOCS[it])]
            if missing_models:
                st.warning("Modèles documentaires manquants pour : " + ", ".join(missing_models))
        st.download_button(
            "⬇️ Télécharger le ZIP",
            data=zip_bytes,
            file_name="Dossier_Audit_ISO42001_{}.zip".format(datetime.now().strftime("%Y%m%d_%H%M")),
            mime="application/zip",
            key="zip_dl_btn"
        )

