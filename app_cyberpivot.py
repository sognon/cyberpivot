# app_cyberpivot.py — CyberPivot™ (FR)
# ============================================================
# - Auth + Admin + Normes Excel (QID/Question -> ID/Contrôle)
# - Édition des contrôles (niveaux FR), commentaires
# - KPI dynamiques (global & vue) : taux pondéré, preuves, etc.
# - UX: "—" si aucun contrôle applicable (évite 0% trompeur)
# - Preuves : upload/list/download/delete + export ZIP (manifest)
# - Exports: DOCX (ISACA), Excel, PDF (si dispo)
# - Thème sombre: valeur KPI visible (contraste corrigé)
# ============================================================

import os, io, json, shutil, hashlib, zipfile, pathlib
from datetime import datetime, timedelta
from typing import Dict, Any, Optional, List

import streamlit as st
import pandas as pd

# Charts
import matplotlib
matplotlib.use("agg")
import matplotlib.pyplot as plt
import numpy as np

# --- Modules internes ---
import auth
import session_guard
import storage
import norms

# ==== Fallback utilitaires (si absents) ====
try:
    import validators  # doit fournir load_norme_excel(file)->DataFrame
except Exception:
    class _V:
        @staticmethod
        def load_norme_excel(f): return pd.read_excel(f)
    validators = _V()

try:
    import errors  # doit fournir report_error(ctx, e)
except Exception:
    class _E:
        @staticmethod
        def report_error(ctx, e): st.warning(f"{ctx} : {e}")
    errors = _E()

# ============================================================
# Config & Styles
# ============================================================
st.set_page_config(page_title="CyberPivot™ – Audit intelligent", page_icon="🛡️", layout="wide")
st.markdown("""
<style>
:root { --brand:#0C2E6B; --brand-soft:#E8EEF9; --accent:#1E40AF; }
.block-container { padding-top: 1rem; }
.badge { display:inline-block; padding:4px 10px; border-radius:999px; background:var(--brand-soft); color:var(--accent); font-weight:700; font-size:12px; }
.headerbar { display:flex; justify-content:space-between; align-items:center; padding:14px 18px; border-radius:14px;
  background:linear-gradient(135deg, #f8fbff 0%, #eef3ff 100%); border:1px solid rgba(16,63,145,0.08);
  box-shadow:0 6px 18px rgba(16,63,145,0.06); margin-bottom:12px; }
.headerbar .title { font-size:20px; font-weight:800; color:var(--brand); }
.kpi { background:#fff; border:1px solid rgba(16,63,145,0.06); border-radius:12px; padding:16px; box-shadow:0 6px 18px rgba(16,63,145,0.06); }
.kpi .value { font-size:28px; font-weight:800; color:#0F172A; } /* <-- contraste forcé (thème sombre) */
.evidence-chip { display:inline-block; padding:3px 8px; border:1px solid #E5E7EB; border-radius:999px; margin:0 6px 6px 0; font-size:12px; }
.evidence-box { border:1px dashed #CBD5E1; padding:12px; border-radius:12px; background:#F8FAFC; }
.small { font-size:12px; color:#6b7280;}
</style>
""", unsafe_allow_html=True)

# ============================================================
# State & Init DB
# ============================================================
def ensure_state():
    for k, v in {
        "std_df": None, "working_df": None, "audit_id": None,
        "_auth_name": None, "_auth_username": None, "_auth_status": False,
        "client_name": "", "contact_name": "", "logo_bytes": None,
        "evidence_map": {},  # key=(ID, Item) -> [ {name, path}, ... ]
    }.items():
        if k not in st.session_state: st.session_state[k] = v
ensure_state()

@st.cache_resource
def _init_all():
    auth.init_auth_db()
    storage.init_db()
    norms.init_norms_db()
    return True
_init_all()

# ============================================================
# Auth
# ============================================================
def _login_ui() -> Dict[str, Any]:
    st.sidebar.subheader("🔐 Authentification")
    tab_co, tab_new = st.sidebar.tabs(["Se connecter", "Créer un compte"])
    with tab_co:
        email = st.text_input("Email", key="login_email")
        pwd = st.text_input("Mot de passe", type="password", key="login_pwd")
        if st.button("Connexion", type="primary"):
            u = auth.verify_password(email, pwd)
            if u and u["is_active"]:
                st.session_state["_auth_name"] = u["full_name"]
                st.session_state["_auth_username"] = u["email"]
                st.session_state["_auth_status"] = True
                st.rerun()
            else:
                st.sidebar.error("Identifiants invalides ou compte inactif.")
    with tab_new:
        n_email = st.text_input("Email (nouveau)")
        n_name = st.text_input("Nom complet")
        n_pwd1 = st.text_input("Mot de passe", type="password")
        n_pwd2 = st.text_input("Confirmer", type="password")
        if st.button("Créer le compte"):
            if not n_email or not n_pwd1: st.sidebar.warning("Email et mot de passe requis.")
            elif n_pwd1 != n_pwd2: st.sidebar.warning("Les mots de passe ne correspondent pas.")
            else:
                if auth.user_exists(n_email): st.sidebar.info("Un compte existe déjà avec cet email.")
                else:
                    auth.create_user(email=n_email, password=n_pwd1, full_name=n_name or n_email,
                                     role="user", tenant_id="default", is_active=True)
                    st.sidebar.success("Compte créé. Connectez-vous.")
                    st.rerun()
    return {
        "name": st.session_state.get("_auth_name"),
        "auth_status": st.session_state.get("_auth_status", False),
        "username": st.session_state.get("_auth_username"),
    }

def _logout_button():
    if st.sidebar.button("Se déconnecter", type="secondary"):
        for k in ["_auth_name","_auth_username","_auth_status"]:
            st.session_state.pop(k, None)
        st.rerun()

auth_ctx = _login_ui()
if not (auth_ctx.get("auth_status") and auth_ctx.get("username")):
    st.stop()

user = session_guard.require_login(
    auth_status=auth_ctx["auth_status"],
    name=auth_ctx["name"],
    username=auth_ctx["username"],
)
USER_EMAIL = user.get("email", auth_ctx["username"])
TENANT_ID  = user.get("tenant_id", "default")
ROLE       = user.get("role", "user")
IS_ADMIN   = (ROLE == "admin")
st.sidebar.success(f"Connecté : {user.get('full_name') or USER_EMAIL} ({ROLE})")
_logout_button()

# ============================================================
# Header
# ============================================================
st.markdown(f"""
<div class="headerbar">
  <div class="title">CyberPivot™ — Audit intelligent</div>
  <div>Connecté : <strong>{auth_ctx.get('name') or 'Utilisateur'}</strong></div>
</div>
""", unsafe_allow_html=True)

# ============================================================
# Helpers (levels, KPI, radar, docx…)
# ============================================================
LEVELS_FR = ["conforme", "partiellement conforme", "non conforme", "non applicable"]
CANON_TO_FR = {
    "yes":"conforme","conforme":"conforme",
    "partial":"partiellement conforme","partially compliant":"partiellement conforme","partiellement conforme":"partiellement conforme",
    "no":"non conforme","non conforme":"non conforme",
    "n/a":"non applicable","na":"non applicable","non applicable":"non applicable","": "non applicable",
}
LEVEL_SCORE = {"conforme":1.0,"partiellement conforme":0.5,"non conforme":0.0,"non applicable":None}
SEVERITY = {"non conforme":"Haut","partiellement conforme":"Moyen","conforme":"Faible","non applicable":"N/A"}

def _to_fr_level(x) -> str:
    if x is None: return "non applicable"
    s = str(x).strip().lower()
    return CANON_TO_FR.get(s, "non applicable")

def _format_kpi(label: str, val: str):
    st.markdown(f"""
    <div class="kpi">
      <div style="font-weight:600;color:#0C2E6B;">{label}</div>
      <div class="value">{val}</div>
    </div>""", unsafe_allow_html=True)

def _compute_metrics(df: pd.DataFrame) -> dict:
    if df is None or df.empty:
        return {"n_total":0,"n_applicable":0,"n_c":0,"n_pc":0,"n_nc":0,"n_na":0,"rate":None}
    d = df.copy()
    d["Level"] = d["Level"].map(_to_fr_level)
    vc = d["Level"].value_counts()
    n_c  = int(vc.get("conforme",0))
    n_pc = int(vc.get("partiellement conforme",0))
    n_nc = int(vc.get("non conforme",0))
    n_na = int(vc.get("non applicable",0))
    n_app = n_c + n_pc + n_nc
    rate = None if n_app==0 else round(((n_c + 0.5*n_pc)/n_app)*100)
    return {"n_total":len(d),"n_applicable":n_app,"n_c":n_c,"n_pc":n_pc,"n_nc":n_nc,"n_na":n_na,"rate":rate}

def _compute_scores(df: pd.DataFrame) -> dict:
    d = df.copy()
    d["Level"] = d["Level"].map(_to_fr_level)
    d["__s"] = d["Level"].map(LEVEL_SCORE)
    d = d[d["__s"].notna()]
    return {
        "global": float(d["__s"].mean()) if not d.empty else 0.0,
        "by_domain": d.groupby("Domain")["__s"].mean().to_dict() if not d.empty else {}
    }

def _radar(scores_by_domain: Dict[str,float]) -> Optional[bytes]:
    if not scores_by_domain: return None
    labels = list(scores_by_domain.keys())
    vals = [max(0.0, min(1.0, scores_by_domain[k])) for k in labels]
    if len(labels) < 3:
        while len(labels) < 3:
            labels.append(labels[-1]+" "); vals.append(vals[-1])
    ang = np.linspace(0, 2*np.pi, len(labels), endpoint=False).tolist()
    vals += vals[:1]; ang += ang[:1]
    fig = plt.figure(figsize=(5,5)); ax = plt.subplot(111, polar=True)
    ax.set_theta_offset(np.pi/2); ax.set_theta_direction(-1)
    ax.set_rlabel_position(0); ax.set_ylim(0,1)
    ax.plot(ang, vals, linewidth=2); ax.fill(ang, vals, alpha=0.25)
    ax.set_xticks(ang[:-1]); ax.set_xticklabels(labels, fontsize=9)
    ax.set_yticks([.25,.5,.75,1]); ax.set_yticklabels(['25%','50%','75%','100%'], fontsize=8)
    bio = io.BytesIO(); fig.tight_layout(); fig.savefig(bio, format="png", dpi=180, bbox_inches="tight"); plt.close(fig)
    return bio.getvalue()

# ---- DOCX minimal (ISACA) ----
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
def _justify_document(doc):
    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs: r.font.size = Pt(11)
def _generate_docx(audit_id: str, df: pd.DataFrame) -> bytes:
    from docx import Document
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(2); s.right_margin = Cm(2)

    title = doc.add_paragraph()
    r = title.add_run(f"Rapport d'audit (ISACA) — {audit_id}")
    r.bold = True; r.font.size = Pt(22)

    d = df.copy()
    for c in ["Domain","ID","Item","Contrôle","Level","Comment"]:
        if c not in d.columns: d[c] = ""
        d[c] = d[c].astype(str)
    d["Level"] = d["Level"].map(_to_fr_level)
    d["_s"] = d["Level"].map(LEVEL_SCORE)
    dscore = d[d["_s"].notna()]
    score_global = round(dscore["_s"].mean()*100) if not dscore.empty else 0
    counts = d["Level"].value_counts().reindex(LEVELS_FR).fillna(0).astype(int)

    doc.add_paragraph(f"Taux de conformité (pondéré) : {score_global}%")
    doc.add_paragraph(
        f"Répartition : {counts.get('conforme',0)} conformes, "
        f"{counts.get('partiellement conforme',0)} partiellement conformes, "
        f"{counts.get('non conforme',0)} non conformes, "
        f"{counts.get('non applicable',0)} non applicables."
    )
    _justify_document(doc)
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def _converter_available():
    try:
        import docx2pdf; return True
    except Exception:
        return shutil.which("soffice") is not None

def _docx_to_pdf_bytes(docx_bytes: bytes) -> Optional[bytes]:
    import tempfile, subprocess
    tmp = tempfile.mkdtemp(prefix="cp_")
    src = os.path.join(tmp, "r.docx"); pdf = os.path.join(tmp, "r.pdf")
    with open(src,"wb") as f: f.write(docx_bytes)
    try:
        import docx2pdf; docx2pdf.convert(src, pdf); 
        with open(pdf,"rb") as r: return r.read()
    except Exception: pass
    soffice = shutil.which("soffice")
    if soffice:
        try:
            subprocess.run([soffice,"--headless","--convert-to","pdf","--outdir",tmp,src], check=True,
                           stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            with open(pdf,"rb") as r: return r.read()
        except Exception as e: errors.report_error("LibreOffice PDF", e)
    return None

# ============================================================
# Évidence (preuves)
# ============================================================
def _slug(s:str)->str:
    s = (s or "").strip().lower()
    ok = "".join(ch if ch.isalnum() or ch in "-_." else "-" for ch in s)
    while "--" in ok: ok = ok.replace("--","-")
    return ok.strip("-_.")
def _evidence_dir(audit:str, qid:str, item:str)->str:
    p = os.path.join("evidence", audit, f"{qid}__{_slug(item)[:60]}")
    pathlib.Path(p).mkdir(parents=True, exist_ok=True); return p
def _has_evidence(audit:str)->bool:
    root = os.path.join("evidence", audit)
    if not os.path.isdir(root): return False
    for _,_,files in os.walk(root):
        if files: return True
    return False
def _persist_uploads(audit:str, qid:str, item:str, files)->List[Dict[str,str]]:
    saved=[]; tgt=_evidence_dir(audit,qid,item)
    for f in files or []:
        ts = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
        name=f"{ts}__{_slug(f.name)}"
        with open(os.path.join(tgt,name),"wb") as w: w.write(f.read())
        saved.append({"name":name,"path":os.path.join(tgt,name)})
    return _load_existing(audit,qid,item)
def _load_existing(audit:str, qid:str, item:str)->List[Dict[str,str]]:
    p=_evidence_dir(audit,qid,item); out=[]
    for n in sorted(os.listdir(p)):
        fp=os.path.join(p,n)
        if os.path.isfile(fp): out.append({"name":n,"path":fp})
    return out
def _delete_file(path:str)->bool:
    try:
        if os.path.isfile(path): os.remove(path); return True
    except Exception as e: errors.report_error("Suppression preuve", e)
    return False
def _export_evidence_zip(audit_id:str)->Optional[bytes]:
    root = os.path.join("evidence", audit_id)
    if not os.path.isdir(root): return None
    entries=[]; total=0; bio=io.BytesIO()
    with zipfile.ZipFile(bio,"w",compression=zipfile.ZIP_DEFLATED) as z:
        for d,_,fs in os.walk(root):
            for fn in fs:
                full=os.path.join(d,fn); rel=os.path.relpath(full,root)
                try:
                    size=os.path.getsize(full); total+=size
                    entries.append({"path":os.path.join(audit_id,rel).replace("\\","/"),"bytes":int(size)})
                    z.write(full, arcname=os.path.join(audit_id,rel))
                except Exception: pass
        if not entries: return None
        import csv
        buf=io.StringIO(); w=csv.writer(buf, lineterminator="\n")
        w.writerow(["path","bytes"])
        for e in entries: w.writerow([e["path"], e["bytes"]])
        z.writestr("manifest.csv", buf.getvalue())
        z.writestr("manifest.json", json.dumps({"audit_id":audit_id,"total_files":len(entries),"total_bytes":total,"entries":entries}, indent=2))
    return bio.getvalue()
def _evidence_stats(audit_id:str, df:pd.DataFrame)->dict:
    tot=0; by={}
    for r in df[["Domain","ID","Item"]].drop_duplicates().itertuples(index=False):
        p=_evidence_dir(audit_id, str(r.ID), str(r.Item))
        c= len([fn for fn in os.listdir(p) if os.path.isfile(os.path.join(p,fn))]) if os.path.isdir(p) else 0
        tot+=c; by[str(r.Domain)]=by.get(str(r.Domain),0)+c
    return {"total":tot,"by_domain":by}

# ============================================================
# Sidebar — params / normes
# ============================================================
st.sidebar.header("⚙️ Paramètres de l’audit")
default_audit = st.session_state.get("audit_id") or f"audit-{datetime.now().strftime('%Y%m%d')}"
audit_id = st.sidebar.text_input("Identifiant de l’audit", value=default_audit)
st.session_state["audit_id"] = audit_id

st.sidebar.subheader("📚 Norme")
norms_av = norms.list_norms(TENANT_ID)
opt = ["(Choisir)"] + [n["name"] for n in norms_av]
sel_norm = st.sidebar.selectbox("Sélectionner une norme publiée", opt, index=0)
if sel_norm != "(Choisir)":
    df_std = norms.get_norm_df(TENANT_ID, sel_norm)
    if df_std is not None:
        st.session_state["std_df"] = df_std
        st.session_state["working_df"] = df_std.copy()
        st.sidebar.success(f"Norme « {sel_norm} » chargée ✅")
    else:
        st.sidebar.error("Impossible de charger la norme.")

st.sidebar.subheader("📄 Rapport")
st.session_state["client_name"]  = st.sidebar.text_input("Client / entité", value=st.session_state.get("client_name",""))
st.session_state["contact_name"] = st.sidebar.text_input("Contact (optionnel)", value=st.session_state.get("contact_name",""))
logo_up = st.sidebar.file_uploader("Logo (PNG/JPG)", type=["png","jpg","jpeg"])
if logo_up: st.session_state["logo_bytes"] = logo_up.read()

# ============================================================
# Navigation
# ============================================================
pages = ["Audit","Mon compte"] + (["Administration"] if IS_ADMIN else [])
page = st.sidebar.radio("Navigation", options=pages, index=0)

# ============================================================
# Page AUDIT
# ============================================================
if page == "Audit":
    st.title("🛡️ Audit")

    def _get_df_or_default() -> pd.DataFrame:
        df = st.session_state.get("working_df")
        if isinstance(df, pd.DataFrame) and not df.empty: return df.copy()
        base = st.session_state.get("std_df")
        if isinstance(base, pd.DataFrame) and not base.empty:
            st.session_state["working_df"] = base.copy(); return base.copy()
        demo = pd.DataFrame([
            {"Domain":"Gouvernance","ID":"GOV-01","Item":"Politique","Contrôle":"Existe-t-il une politique formalisée ?","Level":"non conforme","Comment":""},
            {"Domain":"Sécurité","ID":"SEC-01","Item":"MFA","Contrôle":"MFA activé sur comptes admin ?","Level":"partiellement conforme","Comment":""},
        ])
        st.session_state["std_df"] = demo.copy()
        st.session_state["working_df"] = demo.copy()
        st.warning("Aucune norme sélectionnée : un exemple est chargé.")
        return demo.copy()

    df_all = _get_df_or_default()

    # === Filtres
    f1, f2, f3 = st.columns([1, 2, 1])
    with f1:
        domains = ["(Tous)"] + sorted(df_all["Domain"].dropna().astype(str).unique().tolist())
        dom_sel = st.selectbox("Filtrer par domaine", options=domains, index=0)
    with f2:
        q = st.text_input("Recherche (ID / Item / Contrôle / Commentaire)")
    with f3:
        only_todo = st.toggle("🔎 À traiter (non / partiellement conformes)", value=False)

    def _apply_filters(df: pd.DataFrame) -> pd.DataFrame:
        d = df.copy()
        d["Level"] = d["Level"].map(_to_fr_level)
        if dom_sel != "(Tous)": d = d[d["Domain"] == dom_sel]
        if q.strip():
            qs = q.lower().strip()
            mask = (d["ID"].str.lower().str.contains(qs) |
                    d["Item"].str.lower().str.contains(qs) |
                    d["Contrôle"].str.lower().str.contains(qs) |
                    d["Comment"].str.lower().str.contains(qs))
            d = d[mask]
        if only_todo:
            d = d[d["Level"].isin(["non conforme","partiellement conforme"])]
        return d.reset_index(drop=True)

    view_df = _apply_filters(df_all)

    # === Normalisation colonnes requises
    REQUIRED = ["Domain","ID","Item","Contrôle","Level","Comment"]
    for c in REQUIRED:
        if c not in view_df.columns: view_df[c] = ""
        view_df[c] = view_df[c].astype("string").fillna("").map(lambda v: str(v) if v is not None else "")
    view_df["Level"] = view_df["Level"].map(_to_fr_level)

    # === KPI dynamiques (GLOBAL = working_df, VUE = view_df)
    M_ALL  = _compute_metrics(st.session_state["working_df"])
    M_VIEW = _compute_metrics(view_df)

    c1, c2, c3, c4 = st.columns(4)
    with c1: _format_kpi("Audit", st.session_state.get("audit_id") or "")
    with c2: _format_kpi("Utilisateur", user.get("full_name","Utilisateur"))
    with c3: _format_kpi("Taux de conformité", "—" if M_ALL["rate"] is None else f"{M_ALL['rate']} %")
    with c4:
        ev_stats = _evidence_stats(st.session_state["audit_id"], st.session_state["working_df"])
        _format_kpi("Preuves jointes", str(ev_stats["total"]))

    st.metric("Taux de conformité (global, pondéré)", "—" if M_ALL["rate"] is None else f"{M_ALL['rate']}%")
    if M_ALL["rate"] is None:
        st.info("Aucun contrôle applicable pour l’instant. Renseigne des niveaux (Conforme / Partiellement conforme / Non conforme) pour activer le KPI.")
    st.caption("Méthodologie : (Conformes + 0,5×Partiellement conformes) / (Conformes + Partiellement conformes + Non conformes). N/A exclus.")
    st.divider()

    # === Éditeur (ÉDITION => KPI live)
    st.subheader("✏️ Réponses aux contrôles")
    edited = st.data_editor(
        view_df[REQUIRED],
        key="editor_controls",
        num_rows="dynamic",
        use_container_width=True,
        disabled=["Domain","ID","Item","Contrôle"],
        column_config={
            "Domain":   st.column_config.TextColumn("Domaine"),
            "ID":       st.column_config.TextColumn("ID"),
            "Item":     st.column_config.TextColumn("Item"),
            "Contrôle": st.column_config.TextColumn("Contrôle"),
            "Level": st.column_config.SelectboxColumn(
                "Niveau", options=LEVELS_FR, default="non conforme", required=True,
                help="Niveau de conformité"),
            "Comment": st.column_config.TextColumn("Commentaire"),
        },
        hide_index=True,
    )

    # === 🔄 Synchroniser les modifications de la vue dans le GLOBAL working_df
    def _merge_back(global_df: pd.DataFrame, edited_view: pd.DataFrame) -> pd.DataFrame:
        g = global_df.copy(); e = edited_view.copy()
        key = ["ID","Item"]
        e_map = e.set_index(key)[["Level","Comment"]]
        g_keyed = g.set_index(key)
        g_keyed.update(e_map)  # met à jour Level/Comment sur clés correspondantes
        g = g_keyed.reset_index()
        return g

    st.session_state["working_df"] = _merge_back(st.session_state["working_df"], edited)

    # === PREUVES
    st.subheader("📎 Preuves")
    l, r = st.columns([2,3])
    with l:
        controls = edited[["Domain","ID","Item"]].drop_duplicates().reset_index(drop=True)
        if controls.empty:
            st.info("Aucun contrôle dans la vue.")
            sel = None
        else:
            labels = [f"{r.Domain} — {r.ID} — {r.Item}" for r in controls.itertuples(index=False)]
            sel = st.selectbox("Contrôle", options=list(range(len(labels))), format_func=lambda i: labels[i], index=0)
        if sel is not None:
            row = controls.iloc[sel]
            dom, qid, item = str(row["Domain"]), str(row["ID"]), str(row["Item"])
            up = st.file_uploader("Ajouter des preuves (multi)", type=["pdf","png","jpg","jpeg","xlsx","csv","docx","pptx","txt","log"], accept_multiple_files=True)
            if st.button("➕ Joindre au contrôle"):
                files = up or []
                if not files: st.warning("Aucun fichier sélectionné.")
                else:
                    all_files = _persist_uploads(audit_id, qid, item, files)
                    st.session_state["evidence_map"][(qid,item)] = all_files
                    st.success(f"{len(files)} fichier(s) ajouté(s)."); st.rerun()
    with r:
        if sel is not None:
            qid = str(controls.iloc[sel]["ID"]); item = str(controls.iloc[sel]["Item"])
            ek = (qid,item)
            files_on_disk = _load_existing(audit_id, qid, item)
            st.session_state["evidence_map"][ek] = files_on_disk
            cur = st.session_state["evidence_map"].get(ek, [])
            st.markdown("<div class='evidence-box'>", unsafe_allow_html=True)
            if not cur:
                st.caption("Aucune preuve jointe.")
            else:
                for f in cur:
                    fname, fpath = f["name"], f["path"]
                    cA,cB,cC = st.columns([3,1,1])
                    with cA: st.markdown(f"<span class='evidence-chip'>📄 {fname}</span>", unsafe_allow_html=True)
                    with cB:
                        try:
                            with open(fpath,"rb") as rb:
                                st.download_button("Télécharger", data=rb.read(), file_name=fname, key=f"dl_{fname}")
                        except Exception: st.caption("(introuvable)")
                    with cC:
                        if st.button("❌ Supprimer", key=f"rm_{fname}"):
                            if _delete_file(fpath):
                                st.success("Supprimé."); st.session_state["evidence_map"][ek] = _load_existing(audit_id, qid, item); st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    st.divider()

    # === Sauvegarde DB
    def _clean(x): return "" if x is None else str(x).strip()
    def _payload(row):
        key=( _clean(row.ID), _clean(row.Item) )
        ev = st.session_state["evidence_map"].get(key, [])
        return {
            "domain": _clean(row.Domain), "qid": _clean(row.ID), "item": _clean(row.Item),
            "question": _clean(row.Contrôle), "level": _to_fr_level(_clean(row.Level)),
            "score": None, "criterion":"", "recommendation":"", "comment": _clean(row.Comment),
            "evidence": [{"name":e["name"],"path":e["path"]} for e in ev],
        }

    if st.button("💾 Sauvegarder toutes les réponses", type="primary"):
        saved=0; failed=[]
        for r in st.session_state["working_df"][REQUIRED].itertuples(index=False):
            p=_payload(r)
            if not p["domain"] or not p["qid"] or not p["item"]: continue
            try:
                storage.upsert_response(audit_id, p); saved+=1
            except Exception as e: failed.append((p["qid"], p["item"], str(e)))
        if saved: st.success(f"✅ {saved} réponse(s) sauvegardée(s).")
        if failed: st.error(f"⚠️ {len(failed)} échec(s). Exemple: {failed[0]}")

    st.divider()

    # === Synthèse & dashboard (vue filtrée) — KPI live
    st.subheader("📊 Synthèse (vue filtrée)")
    MV = _compute_metrics(_apply_filters(st.session_state["working_df"]))
    st.metric("Taux de conformité (vue filtrée)", "—" if MV["rate"] is None else f"{MV['rate']}%")
    cA,cB = st.columns([1,2])
    with cA:
        st.caption(f"Sur {MV['n_applicable']} contrôles applicables (N/A exclus) / {MV['n_total']} au total dans la vue.")
        if MV["rate"] is None: st.info("Aucun contrôle applicable dans la vue actuelle.")
    with cB:
        dom_scores = _compute_scores(_apply_filters(st.session_state["working_df"]))["by_domain"]
        if dom_scores:
            dom_df = pd.DataFrame([{"Domaine":k,"Score (%)":round(v*100)} for k,v in dom_scores.items()]).sort_values("Domaine")
            st.dataframe(dom_df, use_container_width=True, hide_index=True)
        else:
            st.caption("Pas de données par domaine pour la vue.")

    with st.expander("👀 Prévisualisation du rapport"):
        prev = _apply_filters(st.session_state["working_df"])
        counts = prev["Level"].map(_to_fr_level).value_counts().reindex(LEVELS_FR).fillna(0).astype(int)
        n = len(prev); val = counts.values.astype(int)
        g1,g2 = st.columns(2)
        with g1:
            fig, ax = plt.subplots(figsize=(5.2,3.2))
            ax.barh(["Conformes","Part. conformes","Non conformes","Non applicables"], val)
            ax.set_xlabel("Nombre de contrôles"); ax.set_title("Répartition par niveau (vue)")
            for i,v in enumerate(val): ax.text(v+0.2, i, str(int(v)), va="center")
            fig.tight_layout(); st.pyplot(fig)
        with g2:
            rpng = _radar(_compute_scores(prev)["by_domain"])
            if rpng: st.image(rpng, caption="Radar par domaine (vue)")
            else: st.caption("Radar indisponible (pas assez de domaines).")
        st.dataframe(prev[REQUIRED].head(20), use_container_width=True, hide_index=True)

    st.divider()

    # === Exports & livrables (live)
    st.subheader("📦 Exports & livrables")
    export_df = st.session_state["working_df"][REQUIRED].copy().fillna("").astype(str)
    # DOCX
    docx_bytes = _generate_docx(audit_id, export_df)
    c1,c2,c3,c4 = st.columns(4)
    with c1:
        st.download_button("📥 Rapport ISACA (DOCX)", data=docx_bytes,
                           file_name=f"rapport_ISACA_{audit_id}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    # Excel
    with c2:
        bio = io.BytesIO()
        with pd.ExcelWriter(bio, engine="openpyxl") as w:
            export_df.to_excel(w, index=False, sheet_name="Audit")
        st.download_button("📊 Export Excel", data=bio.getvalue(),
                           file_name=f"audit_{audit_id}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    # PDF
    with c3:
        conv = _converter_available()
        st.caption("PDF" + (" ✅" if conv else " ⚠️"))
        pdfb = _docx_to_pdf_bytes(docx_bytes) if conv else None
        if pdfb:
            st.download_button("📄 Export PDF", data=pdfb, file_name=f"rapport_ISACA_{audit_id}.pdf", mime="application/pdf")
        else:
            st.caption("Installe docx2pdf ou LibreOffice pour activer la conversion.")
    # ZIP Preuves
    with c4:
        if _has_evidence(audit_id):
            z = _export_evidence_zip(audit_id)
            if z:
                st.download_button("📦 Exporter les preuves (ZIP)", data=z, file_name=f"evidences_{audit_id}.zip", mime="application/zip")
            else:
                st.caption("Aucune preuve à exporter.")
        else:
            st.caption("Aucune preuve trouvée.")

# ============================================================
# Page MON COMPTE
# ============================================================
if page == "Mon compte":
    st.title("👤 Mon compte")
    u = auth.get_user(USER_EMAIL)
    if not u: st.error("Utilisateur introuvable."); st.stop()
    c1,c2 = st.columns(2)
    with c1:
        st.subheader("Profil")
        full = st.text_input("Nom complet", value=u.get("full_name") or "")
        tenant = st.text_input("Tenant / Organisation", value=u.get("tenant_id") or "default")
        if st.button("💾 Mettre à jour le profil", type="primary"):
            try:
                auth.update_user_profile(USER_EMAIL, full_name=full, tenant_id=tenant)
                st.success("Profil mis à jour."); st.rerun()
            except Exception as e: st.error(e)
    with c2:
        st.subheader("Mot de passe")
        p1 = st.text_input("Nouveau mot de passe", type="password")
        p2 = st.text_input("Confirmer", type="password")
        if st.button("🔑 Changer le mot de passe"):
            if not p1 or len(p1)<6: st.warning("Mot de passe trop court.")
            elif p1!=p2: st.warning("Les mots de passe ne correspondent pas.")
            else:
                try: auth.set_password(USER_EMAIL, p1); st.success("Mot de passe mis à jour.")
                except Exception as e: st.error(e)
    st.divider()
    st.json({"email":u["email"],"role":u["role"],"tenant":u["tenant_id"],"actif":u["is_active"],"créé_le":u["created_at"]})

# ============================================================
# Page ADMIN
# ============================================================
if page == "Administration" and IS_ADMIN:
    st.title("🛡️ Administration")
    st.subheader("Utilisateurs")
    try: users = auth.list_users()
    except Exception as e: st.error(e); users=[]
    st.dataframe(pd.DataFrame(users), use_container_width=True) if users else st.caption("Aucun utilisateur.")
    st.subheader("Créer un utilisateur")
    a,b = st.columns(2)
    with a:
        e = st.text_input("Email"); n = st.text_input("Nom complet")
        r = st.selectbox("Rôle", ["user","auditor","admin"]); t = st.text_input("Tenant", value="default")
    with b:
        p = st.text_input("Mot de passe (optionnel)", type="password"); act = st.checkbox("Actif", value=True)
    if st.button("➕ Créer", type="primary"):
        try:
            auth.create_user(email=e, password=p or None, full_name=n, role=r, tenant_id=t, is_active=act)
            st.success("Utilisateur créé."); st.rerun()
        except Exception as ex: st.error(ex)
    st.subheader("Bibliothèque de normes")
    upl = st.file_uploader("Fichier Excel norme", type=["xlsx","xls"])
    name = st.text_input("Nom public de la norme")
    if st.button("📤 Publier / Mettre à jour", type="primary"):
        if not upl or not name.strip(): st.warning("Fichier et nom requis.")
        else:
            try:
                df_norm = validators.load_norme_excel(upl)
                df_norm = df_norm.rename(columns={"QID":"ID","Question":"Contrôle"})
                info = norms.save_norm(TENANT_ID, name.strip(), df_norm)
                st.success(f"Norme publiée : {info['name']}"); st.rerun()
            except Exception as e: st.error(e)
    st.write("### Normes disponibles")
    lst = norms.list_norms(TENANT_ID)
    if lst:
        st.dataframe(pd.DataFrame(lst), use_container_width=True, hide_index=True)
        deln = st.selectbox("Supprimer une norme", ["(Aucune)"]+[n["name"] for n in lst])
        if st.button("🗑️ Supprimer la norme"):
            if deln and deln!="(Aucune)":
                try:
                    if norms.delete_norm(TENANT_ID, deln): st.success("Supprimée."); st.rerun()
                except Exception as e: st.error(e)
    else:
        st.caption("Aucune norme publiée.")

