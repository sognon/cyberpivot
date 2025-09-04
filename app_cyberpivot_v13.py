# app_cyberpivot_v13.py
# -*- coding: utf-8 -*-
from __future__ import annotations
import os, io, re, json, sqlite3
from pathlib import Path
from datetime import date, datetime

import streamlit as st
import pandas as pd
import plotly.graph_objects as go

# Exports
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Local imports
import sys
APP_DIR = Path(__file__).parent.resolve()
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

import auth
import standards

# ========= CONFIG =========
st.set_page_config(page_title="CyberPivot™ — Multi-normes", page_icon="🛡️", layout="wide")
st.markdown("""
<style>
.block-container { max-width: 1200px !important; margin-left:auto; margin-right:auto; padding-top: .6rem; }
.dataframe tr th { text-align:center; background:#0A1F44; color:#fff; }
.dataframe td { text-align:center; }
.smallcap { font-size: 12px; color: #4b5563; }
</style>
""", unsafe_allow_html=True)

# DB & dirs
DB_PATH = os.environ.get("CYBERPIVOT_DB", str(APP_DIR / "cyberpivot.db"))
os.environ["CYBERPIVOT_DB"] = DB_PATH  # aligne auth.py
EVID_DIR = APP_DIR / "evidences"; EVID_DIR.mkdir(parents=True, exist_ok=True)

def db(): return sqlite3.connect(DB_PATH)

def init_app_db():
    with db() as con:
        c = con.cursor()
        c.execute("""CREATE TABLE IF NOT EXISTS projects(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            client TEXT,
            created_at TEXT NOT NULL
        )""")
        c.execute("""CREATE TABLE IF NOT EXISTS audits(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id INTEGER NOT NULL,
            standard TEXT NOT NULL,
            version TEXT,
            created_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(id)
        )""")
        c.execute("""CREATE TABLE IF NOT EXISTS responses(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            audit_id INTEGER NOT NULL,
            domain TEXT NOT NULL,
            qid TEXT NOT NULL,
            question TEXT NOT NULL,
            answer TEXT NOT NULL,
            comment TEXT,
            evidence_json TEXT,
            created_at TEXT NOT NULL,
            UNIQUE(audit_id, qid) ON CONFLICT REPLACE,
            FOREIGN KEY(audit_id) REFERENCES audits(id)
        )""")
        con.commit()

init_app_db()
auth.init_auth_db()

# ========= Risk model =========
DOMAIN_TIERS = {
    "sécurité": ("high", 1.00), "contrôle d’accès": ("high", 1.00),
    "opérations": ("high", 0.90), "sécurité communications": ("high", 0.95),
    "gouvernance ia": ("high", 0.95), "sécurité modèle/données": ("high", 1.00),
    "organisation": ("mid", 0.70), "politiques de sécurité": ("mid", 0.70),
    "gestion des risques ia": ("mid", 0.75), "conformité & éthique": ("mid", 0.75),
    "ressources humaines": ("low", 0.50),
}
KEYWORD_WEIGHTS = [
    (r"\b(mfa|2fa|multi\-?factor)\b",           1.30, 1.20),
    (r"\b(chiffr|encrypt|crypto|tls|ssl)\b",    1.25, 1.20),
    (r"\b(sauvegard|backup|restore)\b",         1.20, 1.15),
    (r"\b(pare\-?feu|firewall|segment|ids|ips)\b", 1.20, 1.15),
    (r"\b(vuln|patch|correctif|cve)\b",         1.20, 1.15),
    (r"\b(journal|log|siem|détect|detect)\b",   1.15, 1.10),
    (r"\b(acc[eè]s|access|privilege|iam)\b",    1.20, 1.15),
    (r"\b(fournisseur|tiers|vendor|third)\b",   1.15, 1.10),
]
STATE = {
    "non conforme":             {"prob": 0.40, "loss_mul": 1.00, "cost_mul": 1.00, "prio": "Haute"},
    "partiellement conforme":   {"prob": 0.22, "loss_mul": 0.75, "cost_mul": 0.75, "prio": "Moyenne"},
    "conforme":                 {"prob": 0.05, "loss_mul": 0.20, "cost_mul": 0.30, "prio": "Info"},
}
BASE_FINANCIALS = {"high": {"loss": 70000, "cost": 17000}, "mid": {"loss": 40000, "cost": 10000}, "low": {"loss": 22000, "cost": 6000}}

def _norm(s: str)->str: return re.sub(r"\s+"," ", (s or "").strip().lower())

def _domain_weight(domain: str)->tuple[str,float]:
    d=_norm(domain)
    for k,(tier,w) in DOMAIN_TIERS.items():
        if k in d: return tier,w
    return "mid",0.70

def _kw_muls(q: str)->tuple[float,float]:
    qn=_norm(q); lm=1.0; cm=1.0
    for pat,l,c in KEYWORD_WEIGHTS:
        if re.search(pat, qn): lm*=l; cm*=c
    return lm,cm

def infer_risk(domain: str, question: str, answer: str)->dict:
    tier,w=_domain_weight(domain)
    base=BASE_FINANCIALS.get(tier, BASE_FINANCIALS["mid"])
    lm,cm=_kw_muls(question)
    stt=STATE.get(_norm(answer), STATE["partiellement conforme"])
    prob=stt["prob"]*w
    loss=base["loss"]*lm*stt["loss_mul"]
    cost=base["cost"]*cm*stt["cost_mul"]
    exp=prob*loss
    return {"probability":prob,"loss_estimate":loss,"remediation_cost":cost,"exp_loss":exp,"priority":stt["prio"]}

def default_reco(answer: str)->str:
    a=_norm(answer)
    if a=="non conforme": return "Mettre en œuvre le contrôle requis et corriger la non-conformité."
    if a=="partiellement conforme": return "Compléter la mise en œuvre jusqu’à conformité totale."
    return "Maintenir, mesurer et documenter la conformité."

def fmt_money(x)->str:
    try: return f"{float(x):,.0f} €".replace(",", " ")
    except: return "—"

# ========= Radar =========
def radar_figure(domain_scores: dict[str,float]) -> go.Figure:
    cats=list(domain_scores.keys()); vals=list(domain_scores.values())
    fig=go.Figure()
    if len(cats)>=3:
        fig.add_trace(go.Scatterpolar(r=vals+[vals[0]], theta=cats+[cats[0]], fill='toself', line_color="#2563EB"))
        fig.update_layout(polar=dict(radialaxis=dict(visible=True,range=[0,100])),
                          showlegend=False, margin=dict(l=40,r=40,t=20,b=20), paper_bgcolor="white")
    return fig

def fig_to_png_bytes(fig: go.Figure, width=900, height=650, scale=2):
    try: return fig.to_image(format="png", width=width, height=height, scale=scale)  # kaleido
    except Exception: return None

# ========= Exports =========
def _docx_header_table(doc, headers):
    t=doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        cell=t.rows[0].cells[i]; cell.text=h
        for run in cell.paragraphs[0].runs:
            run.font.bold=True; run.font.color.rgb=RGBColor(255,255,255)
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'0A1F44'); tcPr.append(shd)
    return t

def export_word(project, standard, audit_id, df, radar_png):
    doc=Document()
    styles=doc.styles
    styles["Normal"].font.name="Segoe UI"; styles["Normal"].font.size=Pt(10)
    for h,sz in [("Heading 1",16),("Heading 2",13),("Heading 3",11)]:
        if h in styles:
            styles[h].font.name="Segoe UI"; styles[h].font.size=Pt(sz); styles[h].font.bold=True; styles[h].font.color.rgb=RGBColor(0x0A,0x1F,0x44)

    # Couverture
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"CyberPivot™\nRapport d’audit — {standard}"); r.bold=True; r.font.size=Pt(24)
    doc.add_paragraph().add_run(f"Projet : {project} • Audit #{audit_id}")
    doc.add_paragraph().add_run(f"Date : {date.today():%d/%m/%Y}")
    doc.add_page_break()

    # TOC
    doc.add_heading("Sommaire", level=1)
    p=doc.add_paragraph(); r=p.add_run()
    fb=OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'),'begin'); r._r.append(fb)
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=r'TOC \o "1-3" \h \z \u'; r._r.append(it)
    fs=OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'),'separate'); r._r.append(fs)
    p.add_run("Table des matières (mise à jour automatique)")
    fe=OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'),'end'); r._r.append(fe)
    doc.add_page_break()

    # KPI
    mapping={"conforme":100,"partiellement conforme":50,"non conforme":0}
    k_total=len(df)
    k_avg=round(df["answer"].str.lower().map(mapping).fillna(0).mean(),1) if k_total else 0.0
    k_conf=(df["answer"].str.lower()=="conforme").sum()
    k_part=(df["answer"].str.lower()=="partiellement conforme").sum()
    k_non =(df["answer"].str.lower()=="non conforme").sum()
    doc.add_heading("1. Résumé exécutif", level=1)
    doc.add_paragraph(f"• Norme : {standard}")
    doc.add_paragraph(f"• Questions évaluées : {k_total}")
    doc.add_paragraph(f"• Conforme : {k_conf} • Partiel : {k_part} • Non conforme : {k_non}")
    doc.add_paragraph(f"• Conformité moyenne : {k_avg}%")
    if radar_png: doc.add_picture(io.BytesIO(radar_png), width=Inches(6.3))
    doc.add_page_break()

    # Constatations
    doc.add_heading("2. Constatations détaillées", level=1)
    t=_docx_header_table(doc, ["Domaine","Contrôle","État","Prob.","Perte (€)","Coût remédiation (€)","Commentaire","Preuves"])
    for _,r in df.iterrows():
        c=t.add_row().cells
        c[0].text=r["domain"]
        c[1].text=f"{r['qid']} — {r['question']}"
        c[2].text=r["answer"]
        c[3].text=f"{r['probability']*100:.0f}%"
        c[4].text=fmt_money(r["loss_estimate"])
        c[5].text=fmt_money(r["remediation_cost"])
        c[6].text=(r.get("comment") or "").strip()
        evs = r.get("evidences", [])
        c[7].text = "\n".join(Path(p).name for p in evs) if evs else "—"

    # Plan d’action
    doc.add_heading("3. Recommandations & plan d’action", level=1)
    pr={"Haute":0,"Moyenne":1,"Basse":2,"Info":3}
    plan=df.sort_values(by=["priority","exp_loss"], key=lambda s: s.map(pr).fillna(9) if s.name=="priority" else -s, ascending=[True,True])
    t2=_docx_header_table(doc, ["Priorité","Domaine","Contrôle","Perte probable (esp.)","Coût remédiation","Recommandation"])
    for _,r in plan.iterrows():
        c=t2.add_row().cells
        c[0].text=r["priority"]; c[1].text=r["domain"]; c[2].text=f"{r['qid']} — {r['question']}"
        c[3].text=fmt_money(r["exp_loss"]); c[4].text=fmt_money(r["remediation_cost"])
        reco = r.get("yaml_recommendation") or default_reco(r["answer"])
        c[5].text=reco

    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf

def export_pdf(project, standard, audit_id, df, radar_png):
    buf=io.BytesIO()
    doc=SimpleDocTemplate(buf, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles=getSampleStyleSheet()
    styles.add(ParagraphStyle("CPH1", parent=styles["Heading1"], textColor=colors.HexColor("#0A1F44")))
    styles.add(ParagraphStyle("CPBody", parent=styles["BodyText"], fontSize=10, leading=13))
    story=[]

    story.append(Paragraph(f"CyberPivot™ — Rapport d’audit — {standard}", styles["Title"]))
    story.append(Spacer(1,6))
    story.append(Paragraph(f"Projet : {project} • Audit #{audit_id}", styles["CPBody"]))
    story.append(Paragraph(f"Date : {date.today():%d/%m/%Y}", styles["CPBody"]))
    story.append(PageBreak())

    story.append(Paragraph("Sommaire", styles["CPH1"]))
    story.append(Paragraph("1. Résumé exécutif", styles["CPBody"]))
    story.append(Paragraph("2. Constatations détaillées", styles["CPBody"]))
    story.append(Paragraph("3. Recommandations & plan d’action", styles["CPBody"]))
    story.append(PageBreak())

    mapping={"conforme":100,"partiellement conforme":50,"non conforme":0}
    k_total=len(df)
    k_avg=round(df["answer"].str.lower().map(mapping).fillna(0).mean(),1) if k_total else 0.0
    k_conf=(df["answer"].str.lower()=="conforme").sum()
    k_part=(df["answer"].str.lower()=="partiellement conforme").sum()
    k_non =(df["answer"].str.lower()=="non conforme").sum()
    story.append(Paragraph("1. Résumé exécutif", styles["CPH1"]))
    story.append(Paragraph(f"Norme : {standard}", styles["CPBody"]))
    story.append(Paragraph(f"Questions évaluées : {k_total}", styles["CPBody"]))
    story.append(Paragraph(f"Conformes : {k_conf} • Partiels : {k_part} • Non conformes : {k_non}", styles["CPBody"]))
    story.append(Paragraph(f"Conformité moyenne : {k_avg}%", styles["CPBody"]))
    story.append(Spacer(1,8))
    if radar_png: story.append(Image(io.BytesIO(radar_png), width=460, height=330))
    story.append(PageBreak())

    # Constatations
    story.append(Paragraph("2. Constatations détaillées", styles["CPH1"]))
    data=[["Domaine","Contrôle","État","Prob.","Perte (€)","Coût remédiation (€)","Commentaire","Preuves"]]
    for _,r in df.iterrows():
        evs = r.get("evidences", [])
        ev_txt = "\n".join(Path(p).name for p in evs) if evs else "—"
        data.append([
            r["domain"], f"{r['qid']} — {r['question']}", r["answer"], f"{r['probability']*100:.0f}%",
            fmt_money(r["loss_estimate"]), fmt_money(r["remediation_cost"]), (r.get("comment") or ""), ev_txt
        ])
    tbl=Table(data, colWidths=[90,220,70,45,75,95,120,120])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0), colors.HexColor("#0A1F44")),
        ("TEXTCOLOR",(0,0),(-1,0), colors.white),
        ("FONTNAME",(0,0),(-1,0), "Helvetica-Bold"),
        ("GRID",(0,0),(-1,-1), 0.4, colors.HexColor("#C8D1DA")),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.whitesmoke, colors.white]),
        ("VALIGN",(0,0),(-1,-1), "TOP"),
        ("ALIGN",(3,1),(5,-1), "CENTER"),
    ]))
    story.append(tbl); story.append(Spacer(1,12))

    # Plan d’action
    story.append(Paragraph("3. Recommandations & plan d’action", styles["CPH1"]))
    pr={"Haute":0,"Moyenne":1,"Basse":2,"Info":3}
    plan=df.sort_values(by=["priority","exp_loss"], key=lambda s: s.map(pr).fillna(9) if s.name=="priority" else -s, ascending=[True,True])
    pdata=[["Priorité","Domaine","Contrôle","Perte probable (esp.)","Coût remédiation","Recommandation"]]
    for _,r in plan.iterrows():
        reco = r.get("yaml_recommendation") or default_reco(r["answer"])
        pdata.append([r["priority"], r["domain"], f"{r['qid']} — {r['question']}",
                      fmt_money(r["exp_loss"]), fmt_money(r["remediation_cost"]), reco])
    pt=Table(pdata, colWidths=[65,90,220,85,95,170])
    pt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0), colors.HexColor("#0A1F44")),
        ("TEXTCOLOR",(0,0),(-1,0), colors.white),
        ("FONTNAME",(0,0),(-1,0), "Helvetica-Bold"),
        ("GRID",(0,0),(-1,-1), 0.4, colors.HexColor("#C8D1DA")),
        ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.whitesmoke, colors.white]),
        ("VALIGN",(0,0),(-1,-1), "TOP"),
    ]))
    story.append(pt)
    doc.build(story); buf.seek(0); return buf

# ========= Auth =========
def login_gate():
    st.sidebar.title("🔐 Authentification")
    with st.sidebar.form("login"):
        email = st.text_input("Email", value="admin@cyberpivot.local")
        pwd = st.text_input("Mot de passe", type="password", value="admin123")
        ok = st.form_submit_button("Se connecter")
    if ok:
        u = auth.verify_password(email, pwd)
        if u:
            st.session_state["user"] = u; st.rerun()
        else:
            st.sidebar.error("Identifiants invalides.")
    return st.session_state.get("user")

if "user" not in st.session_state: st.session_state["user"]=None
user = login_gate()
if not user: st.stop()
if st.sidebar.button("Se déconnecter"): st.session_state["user"]=None; st.rerun()
st.sidebar.success(f"Connecté : {user['full_name']} ({user['role']})")

# ========= Projets & Audits =========
st.title("🛡️ CyberPivot — Multi-normes (Projets • Audits • Rapports)")
st.caption("Questionnaire sans recommandations (uniquement dans les rapports). Preuves par contrôle.")

with st.sidebar.expander("📁 Projets", expanded=True):
    pname = st.text_input("Nom du projet")
    pclient = st.text_input("Client (optionnel)")
    if st.button("➕ Créer projet"):
        if pname:
            with db() as con:
                c=con.cursor()
                c.execute("INSERT OR IGNORE INTO projects(name,client,created_at) VALUES (?,?,?)",
                          (pname.strip(), pclient.strip(), datetime.utcnow().isoformat()))
                con.commit()
            st.success("Projet créé (ou déjà existant)."); st.rerun()
        else:
            st.warning("Nom du projet requis.")
    with db() as con:
        c=con.cursor(); c.execute("SELECT id, name FROM projects ORDER BY created_at DESC")
        plist=c.fetchall()
    pchoices = ["— Aucun —"] + [f"{pid} — {pname}" for pid,pname in plist]
    psel = st.selectbox("Sélectionner un projet", pchoices, index=0, key="sel_project")
    current_project_id = int(psel.split(" — ")[0]) if psel!="— Aucun —" else None

with st.sidebar.expander("🧭 Normes & Audits", expanded=True):
    st.markdown("**Normes disponibles (YAML)**")
    stds = ["— Aucune —"] + standards.list_standards()
    auto_std = st.session_state.get("auto_std")
    idx = stds.index(auto_std) if auto_std in stds else 0
    standard = st.selectbox("Norme", stds, index=idx, key="standard_select")
    ver = st.text_input("Version (ex: 2022 / 2024)", value="")

    if st.button("➕ Créer audit pour ce projet"):
        if not current_project_id:
            st.warning("Sélectionne d’abord un projet.")
        elif standard=="— Aucune —":
            st.warning("Choisis une norme.")
        else:
            with db() as con:
                c=con.cursor()
                c.execute("""INSERT INTO audits(project_id,standard,version,created_at)
                             VALUES (?,?,?,?)""", (current_project_id, standard, ver.strip(), datetime.utcnow().isoformat()))
                con.commit()
            st.success("Audit créé."); st.rerun()

    if current_project_id:
        with db() as con:
            c=con.cursor()
            c.execute("""SELECT id, standard, IFNULL(version,''), created_at
                         FROM audits WHERE project_id=? ORDER BY created_at DESC""", (current_project_id,))
            alist=c.fetchall()
        achoices = ["— Aucun —"] + [f"{aid} — {astd} ({aver})" for aid,astd,aver,ac in alist]
        asel = st.selectbox("Sélectionner un audit", achoices, index=0, key="sel_audit")
        current_audit_id = int(asel.split(" — ")[0]) if asel!="— Aucun —" else None
    else:
        current_audit_id=None

with st.sidebar.expander("📥 Import Excel ⇒ YAML", expanded=False):
    up = st.file_uploader("Excel (FR/EN : Domaine, ID, Question…)", type=["xlsx","xls"])
    std_name_input = st.text_input("Nom de la norme à créer (ex: ISO/IEC 27001)")
    col_imp1, col_imp2 = st.columns(2)
    with col_imp1:
        if st.button("⚙️ Générer le YAML"):
            try:
                if not up or not std_name_input.strip():
                    st.warning("Choisis un Excel et saisis le nom de la norme.")
                else:
                    out = standards.excel_to_yaml(up, std_name_input.strip())
                    st.session_state["auto_std"] = std_name_input.strip()
                    st.success(f"YAML généré : {out}")
                    st.rerun()
            except Exception as e:
                st.error(f"Échec import Excel ⇒ YAML : {e}")
    with col_imp2:
        if st.button("➡️ Utiliser directement ce fichier (sans créer de YAML)"):
            try:
                if not up:
                    st.warning("Sélectionne d’abord un fichier Excel.")
                else:
                    tmp_items = standards.excel_to_items(up)
                    st.session_state["temp_catalog"] = tmp_items
                    st.session_state["temp_std_name"] = std_name_input.strip() or "(Upload temporaire)"
                    st.success(f"{len(tmp_items)} questions chargées depuis l’Excel.")
                    st.rerun()
            except Exception as e:
                st.error(f"Échec lecture directe : {e}")

st.markdown("---")

if not current_project_id:
    st.info("➡️ Sélectionne ou crée un **projet**.")
    st.stop()

if "temp_catalog" in st.session_state:
    catalog = st.session_state["temp_catalog"]
    display_std_name = st.session_state.get("temp_std_name", "Questionnaire (temporaire)")
else:
    if standard == "— Aucune —":
        st.info("➡️ Choisis une **norme** (ou utilise le mode direct depuis Excel).")
        st.stop()
    catalog = standards.load_yaml(standard)
    display_std_name = standard

if not catalog:
    st.error("Aucun item trouvé (YAML ou Excel).")
    st.stop()

if not current_audit_id:
    st.info("➡️ Crée/choisis un **audit** pour ce projet.")
    st.stop()

# Charger preuves existantes
if "evid_map" not in st.session_state:
    st.session_state.evid_map = {}
key_prefix = f"P{current_project_id}:A{current_audit_id}:"
with db() as con:
    c=con.cursor()
    c.execute("SELECT qid, evidence_json FROM responses WHERE audit_id=?", (current_audit_id,))
    for qid, evjson in c.fetchall():
        if evjson:
            try:
                st.session_state.evid_map[key_prefix+qid] = json.loads(evjson)
            except Exception:
                st.session_state.evid_map[key_prefix+qid] = []

# ===== Questionnaire (sans reco) + upload preuves =====
st.subheader(f"📝 Questionnaire — {display_std_name}")
responses=[]
for it in catalog:
    qid = it["id"]; title = it.get("item") or f"{qid} — {it['question']}"
    col1,col2 = st.columns([2.0,1.0])
    with col1:
        st.markdown(f"**{title}**")
        st.caption(it["domain"])
        bits=[]
        if it.get("criterion"): bits.append(f"**Critère :** {it['criterion']}")
        if it.get("objective"): bits.append(f"**Objectif :** {it['objective']}")
        if it.get("reference"): bits.append(f"**Référence :** {it['reference']}")
        if it.get("evidence"):  bits.append(f"**Preuve attendue :** {it['evidence']}")
        if bits: st.write(" / ".join(bits))
        com = st.text_input(f"Commentaire — {qid}", key=f"com_{current_audit_id}_{qid}")
    with col2:
        ans = st.selectbox(f"État — {qid}",
                           ["Conforme","Partiellement conforme","Non conforme"],
                           index=1, key=f"ans_{current_audit_id}_{qid}")
        files = st.file_uploader(f"📎 Preuves — {qid}",
                                 type=["pdf","png","jpg","jpeg","xls","xlsx","csv","txt","docx","pptx","zip"],
                                 accept_multiple_files=True,
                                 key=f"evid_{current_audit_id}_{qid}")
        if st.button(f"Enregistrer les preuves — {qid}", key=f"btn_evid_{current_audit_id}_{qid}"):
            base = (EVID_DIR / f"proj_{current_project_id}" / f"audit_{current_audit_id}" / qid)
            base.mkdir(parents=True, exist_ok=True)
            saved=[]
            for f in files or []:
                name = Path(f.name).name
                dest = base / name
                i=1
                while dest.exists():
                    dest = base / f"{Path(name).stem}_{i}{Path(name).suffix}"
                    i+=1
                with open(dest,"wb") as out: out.write(f.read())
                saved.append(str(dest))
            k = key_prefix+qid
            st.session_state.evid_map[k] = st.session_state.evid_map.get(k, []) + saved
            st.success(f"{len(saved)} fichier(s) enregistré(s)")
        current_evs = st.session_state.evid_map.get(key_prefix+qid, [])
        if current_evs:
            st.caption("Déjà attachées: " + ", ".join(Path(p).name for p in current_evs))

    responses.append({
        "audit_id": current_audit_id,
        "domain": it["domain"], "qid": qid, "question": it["question"],
        "answer": ans, "comment": com,
        "yaml_recommendation": it.get("recommendation",""),
        "evidences": st.session_state.evid_map.get(key_prefix+qid, []),
    })
    st.divider()

if st.button("💾 Enregistrer toutes les réponses", use_container_width=True):
    with db() as con:
        c=con.cursor()
        now=datetime.utcnow().isoformat()
        for r in responses:
            c.execute("""INSERT INTO responses(audit_id,domain,qid,question,answer,comment,evidence_json,created_at)
                         VALUES (?,?,?,?,?,?,?,?)
                         ON CONFLICT(audit_id,qid) DO UPDATE SET
                           domain=excluded.domain,
                           question=excluded.question,
                           answer=excluded.answer,
                           comment=excluded.comment,
                           evidence_json=excluded.evidence_json,
                           created_at=excluded.created_at
                      """, (r["audit_id"], r["domain"], r["qid"], r["question"],
                            r["answer"], r["comment"], json.dumps(r["evidences"]), now))
        con.commit()
    st.success("Réponses enregistrées ✅")
    # Si on utilisait l'Excel direct, on peut le nettoyer après sauvegarde
    if "temp_catalog" in st.session_state:
        st.session_state.pop("temp_catalog", None)
        st.session_state.pop("temp_std_name", None)

if st.button("📊 Calculer & Prévisualiser", use_container_width=True):
    rows=[]
    for r in responses:
        risk=infer_risk(r["domain"], r["question"], r["answer"])
        rows.append({**r, **risk})
    df=pd.DataFrame(rows)

    mapping={"conforme":100,"partiellement conforme":50,"non conforme":0}
    dom_scores = df.assign(score=df["answer"].str.lower().map(mapping).fillna(0)).groupby("domain")["score"].mean().round(1).to_dict()

    st.subheader("📈 Radar des scores par domaine")
    fig=radar_figure(dom_scores)
    st.plotly_chart(fig, use_container_width=True)

    st.subheader("📋 Résultats")
    show=df[["domain","qid","question","answer","comment"]].copy()
    show["preuves"] = df["evidences"].apply(lambda x: ", ".join(Path(p).name for p in x) if isinstance(x,list) else "—")
    st.dataframe(show, use_container_width=True)

    st.markdown("---")
    radar_png = fig_to_png_bytes(fig)

    c1,c2 = st.columns(2)
    with c1:
        if st.button("📝 Générer le rapport Word", use_container_width=True, key="btn_word"):
            buf = export_word(f"{current_project_id}", display_std_name, int(current_audit_id), df, radar_png)
            st.download_button("⬇️ Télécharger Word", data=buf,
                file_name=f"Rapport_{standards._norm_name(display_std_name)}_Audit_{current_audit_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
    with c2:
        if st.button("🧷 Générer le rapport PDF", use_container_width=True, key="btn_pdf"):
            buf = export_pdf(f"{current_project_id}", display_std_name, int(current_audit_id), df, radar_png)
            st.download_button("⬇️ Télécharger PDF", data=buf,
                file_name=f"Rapport_{standards._norm_name(display_std_name)}_Audit_{current_audit_id}.pdf",
                mime="application/pdf", use_container_width=True)

st.caption("CyberPivot™ — Auth • Projets/Audits • Import Excel⇒YAML ou direct • Questionnaire sans reco • Preuves • Radar • Rapports pro.")

