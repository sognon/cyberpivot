# app_cyberpivot_v11_projects_sidebar.py
# ============================================================
# CyberPivot™ — App Streamlit complète (MVP robuste)
# - Login guard (session_guard)
# - Import Excel robuste (validators)
# - Sauvegarde SQLite robuste (storage upsert triple key)
# - Edition des réponses via data_editor
# - Export Word / PDF / Excel
# ============================================================

import os
import io
import json
import time
from datetime import datetime
from typing import List, Dict, Any

import streamlit as st
import pandas as pd

# ---- Modules internes (fourni précédemment) ----
import auth
import session_guard
import errors
import validators
import storage

# Export Word/PDF
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


# ============================================================
# Configuration de page
# ============================================================
st.set_page_config(
    page_title="CyberPivot™ – Audit intelligent",
    page_icon="🛡️",
    layout="wide",
)

# Thème léger (peut être remplacé par tes tokens CSS)
st.markdown("""
<style>
.small-muted { color:#6b7280; font-size:12px; }
.kpi { background:#fff; border:1px solid rgba(16,63,145,0.06); border-radius:12px; padding:16px; box-shadow:0 6px 18px rgba(16,63,145,0.06); }
</style>
""", unsafe_allow_html=True)


# ============================================================
# Initialisation des bases
# ============================================================
@st.cache_resource
def _init_all():
    auth.init_auth_db()
    storage.init_db()
    return True

_init_all()


# ============================================================
# Paramètres
# ============================================================
DEV_MODE = os.getenv("CYBERPIVOT_DEV_MODE", "1") == "1"  # en prod, mets "0"


# ============================================================
# Authentification (exemple : streamlit-authenticator ou autre)
# Ici on simule pour DEV si nécessaire
# ============================================================
def _dev_fake_login() -> Dict[str, Any]:
    """Login fake en DEV pour éviter la config d'auth au démarrage."""
    st.sidebar.info("Mode DEV : connexion simulée")
    return {
        "name": "Dev User",
        "auth_status": True,
        "username": "dev@cyberpivot.local",
    }

def _real_login() -> Dict[str, Any]:
    """
    Adapte ce bloc si tu utilises streamlit-authenticator.
    Doit retourner un dict: {"name":..., "auth_status": bool, "username": ...}
    """
    # Exemple minimaliste : formulaire simple (email + mot de passe)
    st.sidebar.subheader("Se connecter")
    email = st.sidebar.text_input("Email", value="", key="login_email")
    pwd = st.sidebar.text_input("Mot de passe", value="", type="password", key="login_pwd")
    ok = st.sidebar.button("Connexion", type="primary")
    if ok and email and pwd:
        u = auth.verify_password(email, pwd)
        if u:
            st.session_state["_auth_name"] = u["full_name"]
            st.session_state["_auth_username"] = u["email"]
            st.session_state["_auth_status"] = True
            st.sidebar.success("Connecté.")
            # petite pause pour nettoyer l'UI
            time.sleep(0.5)
            st.experimental_rerun()

    return {
        "name": st.session_state.get("_auth_name"),
        "auth_status": st.session_state.get("_auth_status", False),
        "username": st.session_state.get("_auth_username")
    }


if DEV_MODE:
    auth_ctx = _dev_fake_login()
else:
    auth_ctx = _real_login()

# Garde-fou : à partir d’ici, on a toujours un user dict (jamais None)
user = session_guard.require_login(
    auth_status=auth_ctx["auth_status"],
    name=auth_ctx["name"],
    username=auth_ctx["username"],
)

TENANT_ID = user.get("tenant_id", "default")
USER_EMAIL = user.get("email", auth_ctx.get("username") or "user@local")


# ============================================================
# Helpers
# ============================================================
def safe_upsert(audit_id: str, payload: Dict[str, Any]):
    try:
        storage.upsert_response(audit_id, payload)
    except Exception as e:
        errors.report_error("Sauvegarde de la réponse", e)


@st.cache_data(ttl=300)
def _parse_norme_excel(file) -> pd.DataFrame:
    """Parse robuste du fichier de norme."""
    return validators.load_norme_excel(file)


def _format_kpi(label: str, value: str, trend: str | None = None):
    st.markdown(
        f"""
        <div class="kpi">
          <div style="font-weight:600;color:#0C2E6B;">{label}</div>
          <div style="font-size:28px;font-weight:800;">{value}</div>
          {f'<div class="small-muted">{trend}</div>' if trend else ""}
        </div>
        """,
        unsafe_allow_html=True
    )


def _compute_scores(df: pd.DataFrame) -> Dict[str, Any]:
    """Exemple simple de calcul de score : Level -> score 1/0.5/0."""
    map_level = {"Yes": 1.0, "Partial": 0.5, "No": 0.0,
                 "Oui": 1.0, "Partiel": 0.5, "Non": 0.0}
    sc = df.copy()
    sc["__score"] = sc["Level"].map(map_level).fillna(0.0)
    global_score = sc["__score"].mean() if not sc.empty else 0.0
    by_domain = sc.groupby("Domain")["__score"].mean().to_dict() if not sc.empty else {}
    return {"global": global_score, "by_domain": by_domain}


def _generate_word(audit_id: str, df: pd.DataFrame) -> bytes:
    """Génère un DOCX simple : titre + table."""
    document = Document()
    document.add_heading(f"Rapport d'audit – {audit_id}", 0)
    p = document.add_paragraph()
    p.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True

    # Table
    cols = ["Domain", "QID", "Item", "Question", "Level", "Comment"]
    for c in cols:
        if c not in df.columns:
            df[c] = ""

    table = document.add_table(rows=1, cols=len(cols))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = col

    for row in df[cols].itertuples(index=False):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if val is not None else ""

    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()


def _generate_pdf(audit_id: str, df: pd.DataFrame) -> bytes:
    """Génère un PDF texte simple (rapport rapide)."""
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    y = height - 40

    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, f"Rapport d'audit – {audit_id}")
    y -= 20
    c.setFont("Helvetica", 10)
    c.drawString(40, y, f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    y -= 30

    cols = ["Domain", "QID", "Item", "Level"]
    c.setFont("Helvetica-Bold", 10)
    c.drawString(40, y, " | ".join(cols))
    y -= 15
    c.setFont("Helvetica", 10)
    for row in df[cols].itertuples(index=False):
        line = " | ".join([str(v) for v in row])
        if y < 50:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line[:120])
        y -= 12

    c.save()
    return bio.getvalue()


def _generate_excel(audit_id: str, df: pd.DataFrame) -> bytes:
    """Export Excel du travail en cours."""
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Audit")
    return bio.getvalue()


# ============================================================
# Sidebar — Flux principal
# ============================================================
st.sidebar.header("⚙️ Paramètres de l’audit")

# 1) Choix / création d'un audit_id (au minimum)
default_audit_id = st.session_state.get("audit_id") or f"audit-{datetime.now().strftime('%Y%m%d')}"
audit_id = st.sidebar.text_input("Identifiant de l’audit", value=default_audit_id)
st.session_state["audit_id"] = audit_id

# 2) Import de la norme (Excel structuré)
uploaded = st.sidebar.file_uploader("📑 Charger une norme (Excel)", type=["xlsx", "xls"])

if uploaded:
    try:
        df_std = _parse_norme_excel(uploaded)
        st.session_state["std_df"] = df_std.copy()
        st.sidebar.success("Norme chargée ✅")
    except Exception as e:
        errors.report_error("Import du fichier de norme", e)
        st.stop()

# 3) Niveaux de réponse disponibles
LEVELS = ["Yes", "Partial", "No", "N/A"]
level_default = "No"


# ============================================================
# Corps — Dashboard + Edition
# ============================================================
st.title("🛡️ CyberPivot™ – Audit intelligent multinorme")

col1, col2, col3 = st.columns(3)
with col1:
    _format_kpi("Audit", audit_id)
with col2:
    _format_kpi("Utilisateur", user.get("full_name", "Utilisateur"))
with col3:
    _format_kpi("Tenant", user.get("tenant_id", "default"))

st.divider()

# Zone principale : si norme chargée
if "std_df" not in st.session_state:
    st.info("Charge un fichier Excel de norme dans la barre latérale pour commencer.")
    st.stop()

df = st.session_state["std_df"].copy()

# Nettoyage/ajouts de colonnes éditables
for col in ["Level", "Comment", "Score"]:
    if col not in df.columns:
        df[col] = ""

# Préparer une vue "éditable"
edit_df = df[["Domain", "QID", "Item", "Question", "Level", "Comment"]].copy()

# Aide à la saisie : liste déroulante pour Level
# (st.data_editor accepte "column_config" pour suggestions)
level_options = LEVELS

st.subheader("✏️ Réponses aux contrôles")
edited = st.data_editor(
    edit_df,
    num_rows="dynamic",
    use_container_width=True,
    column_config={
        "Level": st.column_config.SelectboxColumn(
            "Level",
            help="Choisis le niveau de conformité",
            options=level_options,
            required=True,
            default=level_default
        ),
        "Comment": st.column_config.TextColumn(
            "Comment",
            help="Commentaire ou précision",
        )
    },
    hide_index=True
)

st.caption("💡 Astuce : tu peux filtrer/ordonner les colonnes avant de sauvegarder.")

# Bouton de sauvegarde
if st.button("💾 Sauvegarder toutes les réponses", type="primary"):
    saved = 0
    for row in edited.itertuples(index=False):
        payload = {
            "domain": getattr(row, "Domain"),
            "qid": getattr(row, "QID"),
            "item": getattr(row, "Item"),
            "question": getattr(row, "Question"),
            "level": getattr(row, "Level") or level_default,
            "score": None,  # si tu veux mapper Level -> score, tu peux calculer ici
            "criterion": None,
            "recommendation": None,
            "comment": getattr(row, "Comment") or "",
            "evidence": [],
        }
        safe_upsert(audit_id, payload)
        saved += 1
    st.success(f"✅ {saved} réponse(s) sauvegardée(s).")

st.divider()

# Synthèse
scores = _compute_scores(edited)
st.subheader("📊 Synthèse")
c1, c2 = st.columns([1, 2])
with c1:
    _format_kpi("Conformité globale", f"{round(scores['global']*100)} %")
with c2:
    if scores["by_domain"]:
        dom_df = pd.DataFrame(
            [{"Domain": k, "Score (%)": round(v*100)} for k, v in scores["by_domain"].items()]
        ).sort_values("Domain")
        st.dataframe(dom_df, use_container_width=True, hide_index=True)
    else:
        st.info("Aucune donnée de score par domaine pour le moment.")

st.divider()

# Exports
st.subheader("📦 Exports & livrables")

colx, coly, colz = st.columns(3)
with colx:
    if st.button("📝 Générer Word (DOCX)"):
        try:
            docx_bytes = _generate_word(audit_id, edited)
            st.success("DOCX généré ✅")
            st.download_button("⬇️ Télécharger le DOCX", data=docx_bytes,
                               file_name=f"rapport_{audit_id}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            errors.report_error("Génération Word", e)

with coly:
    if st.button("📄 Générer PDF"):
        try:
            pdf_bytes = _generate_pdf(audit_id, edited)
            st.success("PDF généré ✅")
            st.download_button("⬇️ Télécharger le PDF", data=pdf_bytes,
                               file_name=f"rapport_{audit_id}.pdf",
                               mime="application/pdf")
        except Exception as e:
            errors.report_error("Génération PDF", e)

with colz:
    if st.button("📊 Export Excel"):
        try:
            xlsx_bytes = _generate_excel(audit_id, edited)
            st.success("Excel généré ✅")
            st.download_button("⬇️ Télécharger l’Excel", data=xlsx_bytes,
                               file_name=f"audit_{audit_id}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        except Exception as e:
            errors.report_error("Export Excel", e)

st.markdown("<div class='small-muted'>© CyberPivot™ — MVP de démonstration</div>", unsafe_allow_html=True)

